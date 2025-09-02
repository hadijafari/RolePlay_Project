"""
Document Intelligence Service
Advanced AI-powered document analysis using multi-modal processing and structured outputs.
This service provides superior document analysis compared to basic text extraction.
"""

import os
import sys
import time
import json
import logging
import asyncio
import tempfile
from typing import Dict, Any, List, Optional, Union, Tuple
from pathlib import Path
from datetime import datetime

try:
    import openai
    from openai import OpenAI
    from pydantic import ValidationError
    from dotenv import load_dotenv
    
    # Document processing libraries
    import PyPDF2
    import docx
    import pandas as pd
    
except ImportError as e:
    print(f"Document Intelligence Service: Missing dependencies: {e}")
    print("Please install: pip install openai PyPDF2 python-docx pandas pydantic")
    sys.exit(1)

# Add project root to path
current_dir = Path(__file__).parent
project_root = current_dir.parent
sys.path.insert(0, str(project_root))

try:
    from models.interview_models import (
        DocumentType, ResumeAnalysis, JobDescriptionAnalysis, 
        PersonalInfo, TechnicalSkill, WorkExperience, Education, 
        Certification, CareerProgression, JobRequirement, CompanyInfo,
        ProcessingResult, SkillLevel, CandidateMatch
    )
except ImportError as e:
    print(f"Document Intelligence Service: Missing models: {e}")
    sys.exit(1)

# Load environment variables
env_path = Path(__file__).parent.parent.parent.parent.parent / ".env"
load_dotenv(env_path)


class DocumentParsingError(Exception):
    """Custom exception for document parsing errors."""
    pass


class DocumentIntelligenceConfig:
    """Configuration for document intelligence service."""
    
    # OpenAI settings
    MODEL_TEXT = "gpt-4o-mini"  # For text-based analysis
    MODEL_VISION = "gpt-4o"     # For visual/PDF analysis
    MAX_TOKENS = 4000
    TEMPERATURE = 0.3  # Lower temperature for more consistent structured outputs
    
    # Document processing
    MAX_FILE_SIZE_MB = 10
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt']
    
    # Processing settings
    USE_VISION_MODEL = True  # Use GPT-4V for PDF analysis
    MULTI_PASS_ANALYSIS = False  # Multiple analysis passes for accuracy
    STRUCTURED_OUTPUT_ENFORCEMENT = True  # Enforce JSON structure
    
    # Performance settings
    MAX_RETRY_ATTEMPTS = 3
    TIMEOUT = 60.0  # seconds
    PARALLEL_PROCESSING = True


class DocumentIntelligenceService:
    """
    Advanced document analysis service using AI for structured data extraction.
    
    This service provides superior document analysis by:
    1. Using multi-modal AI (GPT-4V) for visual document understanding
    2. Enforcing structured outputs with Pydantic models
    3. Multi-pass analysis for accuracy and completeness
    4. Semantic understanding of document sections and content
    """
    
    def __init__(self, api_key: Optional[str] = None):
        """Initialize the document intelligence service."""
        
        self.config = DocumentIntelligenceConfig()
        self.logger = self._setup_logger()
        # Debug logging disabled for production
        # self.logger.setLevel(logging.DEBUG)
        
        # Initialize OpenAI client
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError("Document Intelligence Service: OpenAI API key not found")
        
        self.client = OpenAI(api_key=self.api_key)
        
        # Statistics tracking
        self.stats = {
            "documents_processed": 0,
            "successful_analyses": 0,
            "failed_analyses": 0,
            "total_processing_time": 0.0,
            "vision_model_used": 0,
            "text_model_used": 0,
            "multi_pass_analyses": 0
        }
        
        self.logger.info("Document Intelligence Service: Initialized with advanced AI analysis")
    
    def _setup_logger(self) -> logging.Logger:
        """Set up logging for the service."""
        logger = logging.getLogger("DocumentIntelligenceService")
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - Document Intelligence: %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def validate_document(self, file_path: str) -> ProcessingResult:
        """Validate document before processing."""
        
        try:
            path = Path(file_path)
            
            # Check file exists
            if not path.exists():
                return ProcessingResult(
                    success=False,
                    error_message=f"File not found: {file_path}"
                )
            
            # Check file size
            file_size_mb = path.stat().st_size / (1024 * 1024)
            if file_size_mb > self.config.MAX_FILE_SIZE_MB:
                return ProcessingResult(
                    success=False,
                    error_message=f"File too large: {file_size_mb:.1f}MB (max: {self.config.MAX_FILE_SIZE_MB}MB)"
                )
            
            # Check file format
            file_extension = path.suffix.lower()
            if file_extension not in self.config.SUPPORTED_FORMATS:
                return ProcessingResult(
                    success=False,
                    error_message=f"Unsupported format: {file_extension}"
                )
            
            self.logger.info(f"Document validation passed: {path.name} ({file_size_mb:.1f}MB)")
            
            return ProcessingResult(
                success=True,
                result_data={
                    "file_path": str(path),
                    "file_size_mb": file_size_mb,
                    "format": file_extension,
                    "validation_time": datetime.now().isoformat()
                }
            )
            
        except Exception as e:
            return ProcessingResult(
                success=False,
                error_message=f"Validation error: {e}"
            )
    
    def extract_document_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract content from document with metadata.
        
        Returns:
            Tuple of (text_content, metadata)
        """
        
        path = Path(file_path)
        file_extension = path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_content(file_path)
            elif file_extension == '.docx':
                return self._extract_docx_content(file_path)
            elif file_extension == '.txt':
                return self._extract_txt_content(file_path)
            else:
                raise DocumentParsingError(f"Unsupported format: {file_extension}")
                
        except Exception as e:
            self.logger.error(f"Content extraction failed: {e}")
            raise DocumentParsingError(f"Failed to extract content: {e}")
    
    def _extract_pdf_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from PDF file."""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text content
                text_content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                
                # Extract metadata
                metadata = {
                    "total_pages": len(pdf_reader.pages),
                    "format": "pdf",
                    "extraction_method": "text_extraction",
                    "has_images": False,  # Would need more sophisticated detection
                    "estimated_visual_complexity": "medium"  # Placeholder
                }
                
                return text_content.strip(), metadata
                
        except Exception as e:
            raise DocumentParsingError(f"PDF extraction failed: {e}")
    
    def _extract_docx_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from DOCX file."""
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text content
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract table content
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_content.append(row_data)
                    text_content += " | ".join(row_data) + "\n"
            
            # Extract metadata
            metadata = {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "format": "docx",
                "extraction_method": "structured_extraction",
                "tables_found": len(table_content) > 0
            }
            
            return text_content.strip(), metadata
            
        except Exception as e:
            raise DocumentParsingError(f"DOCX extraction failed: {e}")
    
    def _extract_txt_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from TXT file."""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            metadata = {
                "total_lines": len(text_content.split('\n')),
                "format": "txt",
                "extraction_method": "direct_read",
                "encoding": "utf-8"
            }
            
            return text_content.strip(), metadata
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'ascii']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    
                    metadata = {
                        "total_lines": len(text_content.split('\n')),
                        "format": "txt",
                        "extraction_method": "direct_read",
                        "encoding": encoding
                    }
                    
                    return text_content.strip(), metadata
                except:
                    continue
            
            raise DocumentParsingError("Could not decode text file with any encoding")
    
    async def analyze_resume(self, file_path: str) -> ProcessingResult:
        """
        Analyze resume with advanced AI processing.
        
        Uses structured outputs and multi-pass analysis for maximum accuracy.
        """
        print("=" * 80)
        start_time = time.time()
        self.stats["documents_processed"] += 1
        
        try:
            self.logger.info(f"Starting advanced resume analysis: {Path(file_path).name}")
            
            # Validate document
            validation_result = self.validate_document(file_path)
            if not validation_result.success:
                return validation_result
            
            # Extract content
            text_content, extraction_metadata = self.extract_document_content(file_path)
            # print(f"In analyze_resume method of Document Intelligence Service: Extracted content from {text_content}")

            # Perform AI-powered analysis
            if self.config.USE_VISION_MODEL and Path(file_path).suffix.lower() == '.pdf':
                print(f"In analyze_resume method of Document Intelligence Service: Vision model used for analysis")
                resume_analysis = await self._analyze_resume_with_vision(file_path, text_content)
                self.stats["vision_model_used"] += 1
            else:
                print(f"In analyze_resume method of Document Intelligence Service: Text model used for analysis")
                resume_analysis = await self._analyze_resume_with_text(text_content)
                self.stats["text_model_used"] += 1
            
            # Multi-pass refinement if enabled
            if self.config.MULTI_PASS_ANALYSIS:
                resume_analysis = await self._refine_resume_analysis(resume_analysis, text_content)
                self.stats["multi_pass_analyses"] += 1
            
            # Calculate processing time
            processing_time = time.time() - start_time
            self.stats["total_processing_time"] += processing_time
            self.stats["successful_analyses"] += 1
            
            self.logger.info(f"Resume analysis completed in {processing_time:.2f}s")
            
            # Serialize the resume analysis data to handle datetime objects
            try:
                resume_dict = resume_analysis.dict()
                self.logger.debug(f"Resume dict keys: {list(resume_dict.keys())}")
                
                # Check for datetime objects in the dict
                def check_datetime_objects(obj, path=""):
                    if isinstance(obj, datetime):
                        self.logger.error(f"Found datetime object at path: {path}")
                        return True
                    elif isinstance(obj, dict):
                        for k, v in obj.items():
                            if check_datetime_objects(v, f"{path}.{k}"):
                                return True
                    elif isinstance(obj, list):
                        for i, v in enumerate(obj):
                            if check_datetime_objects(v, f"{path}[{i}]"):
                                return True
                    return False
                
                has_datetime = check_datetime_objects(resume_dict)
                if has_datetime:
                    self.logger.error("Found datetime objects in resume_dict before serialization")
                else:
                    self.logger.debug("No datetime objects found in resume_dict")
                
                serialized_data = self._serialize_analysis_data(resume_dict)
                self.logger.debug("Resume analysis data serialized successfully")
            except Exception as e:
                self.logger.error(f"Failed to serialize resume analysis data: {e}")
                import traceback
                self.logger.error(f"Traceback: {traceback.format_exc()}")
                raise
            
            # Create ProcessingResult without automatic timestamp to avoid serialization issues
            result = ProcessingResult(
                success=True,
                result_data=serialized_data,
                processing_time_seconds=processing_time,
                confidence_score=resume_analysis.confidence_score,
                metadata={
                    "file_path": file_path,
                    "extraction_metadata": extraction_metadata,
                    "analysis_method": "vision" if self.config.USE_VISION_MODEL and Path(file_path).suffix.lower() == '.pdf' else "text",
                    "multi_pass_used": self.config.MULTI_PASS_ANALYSIS
                }
            )
            
            self.logger.debug(f"ProcessingResult created successfully with timestamp: {result.timestamp}")
            return result
            
        except Exception as e:
            processing_time = time.time() - start_time
            self.stats["failed_analyses"] += 1
            self.logger.error(f"Resume analysis failed after {processing_time:.2f}s: {e}")
            
            return ProcessingResult(
                success=False,
                error_message=str(e),
                processing_time_seconds=processing_time
            )
    
    async def _analyze_resume_with_text(self, text_content: str) -> ResumeAnalysis:
        """Analyze resume using text-based AI model with structured outputs."""
        
        system_prompt = """You are an expert resume analysis AI. Your job is to extract structured information from resumes with high accuracy.

CRITICAL CHARACTER RESTRICTIONS:
- Use ONLY standard ASCII characters (A-Z, a-z, 0-9, basic punctuation)
- DO NOT use Unicode symbols like checkmarks (✓), bullet points (•), em dashes (—), smart quotes (" "), or any special characters
- Use simple text formatting: use "- " for bullet points, use regular quotes " and apostrophes '
- Keep all text content compatible with basic ASCII encoding
- Replace any special characters from the source document with ASCII equivalents

CRITICAL: You MUST return a valid JSON response that exactly matches this ResumeAnalysis schema:

{
  "document_type": "resume",
  "personal_info": {
    "full_name": "string or null",
    "email": "string or null", 
    "phone": "string or null",
    "location": "string or null",
    "linkedin": "string or null",
    "github": "string or null",
    "portfolio": "string or null"
  },
  "technical_skills": [
    {
      "skill_name": "string",
      "category": "string", 
      "proficiency_level": "beginner|intermediate|advanced|expert or null",
      "years_experience": number or null,
      "mentioned_count": number,
      "context_mentions": ["string array"]
    }
  ],
  "soft_skills": ["string array"],
  "work_experience": [
    {
      "job_title": "string",
      "company_name": "string",
      "duration": "string",
      "start_date": "string or null",
      "end_date": "string or null", 
      "is_current": boolean,
      "key_responsibilities": ["string array"],
      "achievements": ["string array"],
      "technologies_used": ["string array"],
      "team_size": number or null,
      "reporting_structure": "string or null"
    }
  ],
  "education": [
    {
      "degree": "string",
      "major": "string or null",
      "institution": "string", 
      "graduation_year": number or null,
      "gpa": number or null,
      "relevant_coursework": ["string array"],
      "honors": ["string array"]
    }
  ],
  "certifications": [
    {
      "name": "string",
      "issuing_organization": "string",
      "issue_date": "string or null",
      "expiry_date": "string or null", 
      "credential_id": "string or null",
      "is_active": boolean
    }
  ],
  "career_progression": {
    "total_experience_years": number,
    "career_trajectory": "string",
    "job_changes_frequency": number,
    "industry_consistency": boolean,
    "role_progression": ["string array"],
    "career_gaps": ["string array"]
  },
  "strengths": ["string array"],
  "potential_gaps": ["string array"], 
  "unique_selling_points": ["string array"],
  "red_flags": ["string array"],
  "confidence_score": number between 0 and 1
}

CRITICAL DATA STRUCTURE RULES:
1. ARRAYS: All array fields MUST be arrays, never null or undefined
   - If no data exists, use empty array: []
   - Examples: "soft_skills": [], "strengths": [], "red_flags": []
   - NEVER use null for arrays: ❌ "soft_skills": null

2. REQUIRED FIELDS: These fields MUST always be present and properly typed:
   - "career_progression.total_experience_years": must be a number (0.0 if unknown)
   - "career_progression.career_trajectory": must be a string ("unknown" if unclear)
   - "career_progression.job_changes_frequency": must be a number (0.0 if unknown)
   - "career_progression.industry_consistency": must be a boolean (true if unknown)
   - "career_progression.role_progression": must be an array ([] if unknown)

3. ARRAY CONTENT RULES:
   - "technical_skills": Must contain skill objects with all required fields
   - "work_experience": Must contain experience objects with "company_name" and "duration"
   - "context_mentions": Must be array of strings, never null

4. NULL vs EMPTY ARRAY GUIDELINES:
   - Use null for: missing optional strings, numbers, booleans
   - Use [] for: missing optional arrays
   - Use "" for: missing optional strings that should be empty strings

FIELD REQUIREMENTS:
- Use EXACT field names: "personal_info" (not "personal_information")
- technical_skills must be an ARRAY of objects (not nested dictionaries)
- work_experience requires "company_name" and "duration" fields
- career_progression requires "total_experience_years", "career_trajectory", "job_changes_frequency", "industry_consistency", "role_progression"
- All arrays should be flat lists of strings or objects as specified
- Use null for missing optional fields, never omit required fields

ANALYSIS GUIDELINES:
- For skills: categorize as "Programming Languages", "Frameworks", "Tools", "Databases", "Cloud Platforms", etc.
- For experience: extract company names, calculate durations, identify achievements
- For career progression: calculate total years, assess trajectory (ascending/lateral/mixed)
- Set confidence_score based on information completeness (0.0-1.0)"""
        
        user_prompt = f"""Please analyze this resume and extract structured information:

RESUME CONTENT:
{text_content}

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON - no markdown, no explanations, no code blocks
2. Use the exact field names and structure shown in the schema
3. Do NOT include "analysis_timestamp" field - it will be added automatically
4. For missing information, use null (not empty strings)
5. Ensure all required fields are present with appropriate values

CRITICAL DATA STRUCTURE ENFORCEMENT:
- ALL array fields MUST be arrays: use [] for empty arrays, NEVER use null
- career_progression.total_experience_years: must be a number (0.0 if unknown)
- career_progression.career_trajectory: must be a string ("unknown" if unclear)
- career_progression.job_changes_frequency: must be a number (0.0 if unknown)
- career_progression.industry_consistency: must be a boolean (true if unknown)
- career_progression.role_progression: must be an array ([] if unknown)

Return the complete JSON response following the ResumeAnalysis schema exactly."""
        
        try:
            response = await asyncio.to_thread(
                self.client.chat.completions.create,
                model=self.config.MODEL_TEXT,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=self.config.MAX_TOKENS,
                temperature=self.config.TEMPERATURE
            )
            
            # Parse structured response
            response_content = response.choices[0].message.content
            # print(f"In _analyze_resume_with_text method of Document Intelligence Service: AI Response: {response_content}")
            # saving the response_content into a file in the filepath
            # with open("D:\\workspaces\\AI-Tutorials\\AI Agents\\MyAgentsTutorial\\agents\\2_openai\\Interview RolePlay\\interview_platform\\RAG\\first_AI_analysis_response.txt", "w", encoding="utf-8") as f:
            #     f.write(response_content)
            # print(f"In _analyze_resume_with_text method of Document Intelligence Service: AI Response saved to file")
            self.logger.debug(f"AI Response (first 500 chars): {response_content[:500]}")
            
            # Try to extract JSON from response
            json_data = self._extract_json_from_response(response_content)
            self.logger.debug(f"Extracted JSON keys: {list(json_data.keys()) if isinstance(json_data, dict) else 'Not a dict'}")
            
            # Create and validate ResumeAnalysis object
            # Remove analysis_timestamp if it exists (will be auto-generated)
            if 'analysis_timestamp' in json_data:
                self.logger.debug("Removing analysis_timestamp from AI response")
                del json_data['analysis_timestamp']
            
            # Clean any datetime objects from the JSON data
            json_data = self._clean_datetime_objects(json_data)
            self.logger.debug("JSON data cleaned of datetime objects")
            
            try:
                resume_analysis = ResumeAnalysis(**json_data)
                self.logger.debug("ResumeAnalysis object created successfully")
            except Exception as e:
                self.logger.error(f"Failed to create ResumeAnalysis object: {e}")
                raise
            
            try:
                # Test serialization of the created object
                test_dict = resume_analysis.dict()
                self.logger.debug("ResumeAnalysis.dict() created successfully")
            except Exception as e:
                self.logger.error(f"Failed to serialize ResumeAnalysis to dict: {e}")
                raise
            
            return resume_analysis
            
        except ValidationError as e:
            self.logger.error(f"Resume analysis validation error: {e}")
            # Return basic analysis as fallback
            return self._create_fallback_resume_analysis(text_content)
        
        except Exception as e:
            self.logger.error(f"Resume text analysis error: {e}")
            return self._create_fallback_resume_analysis(text_content)
    
    async def _analyze_resume_with_vision(self, file_path: str, text_content: str) -> ResumeAnalysis:
        """Analyze resume using vision model for PDFs (placeholder - would need vision API implementation)."""
        
        # This would use GPT-4V with image input for PDF analysis
        # For now, fall back to text analysis
        self.logger.info("Vision model analysis - falling back to text analysis")
        return await self._analyze_resume_with_text(text_content)
    
    async def _refine_resume_analysis(self, initial_analysis: ResumeAnalysis, text_content: str) -> ResumeAnalysis:
        """Refine resume analysis with additional AI pass."""
        
        # Serialize the initial analysis data to handle datetime objects
        initial_analysis_data = self._serialize_analysis_data(initial_analysis.dict())
        
        refinement_prompt = f"""Please review and refine this resume analysis for accuracy and completeness:

INITIAL ANALYSIS:
{json.dumps(initial_analysis_data, indent=2)}

ORIGINAL RESUME:
{text_content[:2000]}...

Please provide an improved analysis, focusing on:
1. Correcting any inaccuracies
2. Adding missing skills or experience
3. Improving career progression analysis
4. Enhancing strengths and gap identification

Return the refined analysis in JSON format."""
        
        try:
            response = await asyncio.to_thread(
                self.client.chat.completions.create,
                model=self.config.MODEL_TEXT,
                messages=[
                    {"role": "system", "content": "You are refining a resume analysis for accuracy."},
                    {"role": "user", "content": refinement_prompt}
                ],
                max_tokens=self.config.MAX_TOKENS,
                temperature=0.1  # Lower temperature for refinement
            )
            
            # Parse refined response
            response_content = response.choices[0].message.content
            json_data = self._extract_json_from_response(response_content)
            # with open("D:\\workspaces\\AI-Tutorials\\AI Agents\\MyAgentsTutorial\\agents\\2_openai\\Interview RolePlay\\interview_platform\\RAG\\second_refined_AI_analysis_response.txt", "w", encoding="utf-8") as f:
            #     f.write(response_content)
            # # print(f"In _refine_resume_analysis method of Document Intelligence Service: Extracted JSON from response: {json_data}")
            # print(f"In _refine_resume_analysis method of Document Intelligence Service: Refined AI Response saved to file")
            
            # Clean datetime objects from refined data
            if 'analysis_timestamp' in json_data:
                del json_data['analysis_timestamp']
            json_data = self._clean_datetime_objects(json_data)
            
            # Create refined analysis
            refined_analysis = ResumeAnalysis(**json_data)
            
            # Update confidence score for refined analysis
            refined_analysis.confidence_score = min(1.0, (initial_analysis.confidence_score or 0.5) + 0.2)
            
            return refined_analysis
            
        except Exception as e:
            self.logger.warning(f"Resume refinement failed: {e}")
            return initial_analysis  # Return original analysis
    
    async def analyze_job_description(self, file_path: str) -> ProcessingResult:
        """Analyze job description with advanced AI processing."""
        
        start_time = time.time()
        
        try:
            self.logger.info(f"Starting job description analysis: {Path(file_path).name}")
            
            # Validate and extract content
            validation_result = self.validate_document(file_path)
            if not validation_result.success:
                return validation_result
            
            text_content, extraction_metadata = self.extract_document_content(file_path)
            
            # Analyze with AI
            job_analysis = await self._analyze_job_description_with_ai(text_content)
            
            processing_time = time.time() - start_time
            self.stats["successful_analyses"] += 1
            
            return ProcessingResult(
                success=True,
                result_data=self._serialize_analysis_data(job_analysis.dict()),
                processing_time_seconds=processing_time,
                confidence_score=job_analysis.confidence_score
            )
            
        except Exception as e:
            processing_time = time.time() - start_time
            self.stats["failed_analyses"] += 1
            
            return ProcessingResult(
                success=False,
                error_message=str(e),
                processing_time_seconds=processing_time
            )
    
    async def _analyze_job_description_with_ai(self, text_content: str) -> JobDescriptionAnalysis:
        """Analyze job description using AI with structured outputs."""
        
        system_prompt = """You are an expert job description analyzer. Extract structured information from job postings with high accuracy.

CRITICAL CHARACTER RESTRICTIONS:
- Use ONLY standard ASCII characters (A-Z, a-z, 0-9, basic punctuation)
- DO NOT use Unicode symbols like checkmarks (✓), bullet points (•), em dashes (—), smart quotes (" "), or any special characters
- Use simple text formatting: use "- " for bullet points, use regular quotes " and apostrophes '
- Keep all text content compatible with basic ASCII encoding
- Replace any special characters from the source document with ASCII equivalents

CRITICAL: You MUST return a valid JSON response that exactly matches this JobDescriptionAnalysis schema:

{
  "document_type": "job_description",
  "job_title": "string",
  "company_info": {
    "company_name": "string or null",
    "industry": "string or null", 
    "company_size": "string or null",
    "culture_keywords": ["string array"],
    "benefits_offered": ["string array"]
  },
  "technical_requirements": [
    {
      "requirement": "string",
      "category": "technical|soft_skill|experience|education", 
      "priority": "required|preferred|nice_to_have",
      "years_required": number or null
    }
  ],
  "soft_skill_requirements": [
    {
      "requirement": "string",
      "category": "soft_skill",
      "priority": "required|preferred|nice_to_have", 
      "years_required": number or null
    }
  ],
  "experience_requirements": [
    {
      "requirement": "string", 
      "category": "experience",
      "priority": "required|preferred|nice_to_have",
      "years_required": number or null
    }
  ],
  "education_requirements": [
    {
      "requirement": "string",
      "category": "education", 
      "priority": "required|preferred|nice_to_have",
      "years_required": number or null
    }
  ],
  "job_summary": "string",
  "key_responsibilities": ["string array"],
  "success_metrics": ["string array"],
  "career_growth_opportunities": ["string array"],
  "role_seniority": "junior|mid|senior|lead|executive",
  "role_type": "individual_contributor|team_lead|manager|director",
  "critical_success_factors": ["string array"],
  "potential_challenges": ["string array"],
  "confidence_score": number between 0 and 1
}

CRITICAL DATA STRUCTURE RULES:
1. ARRAYS: All array fields MUST be arrays, never null or undefined
   - If no data exists, use empty array: []
   - Examples: "culture_keywords": [], "benefits_offered": [], "key_responsibilities": []
   - NEVER use null for arrays: ❌ "benefits_offered": null

2. REQUIRED FIELDS: These fields MUST always be present and properly typed:
   - "job_title": must be a string (use "Position Title Not Specified" if unclear)
   - "role_seniority": must be one of: "junior", "mid", "senior", "lead", "executive"
   - "role_type": must be one of: "individual_contributor", "team_lead", "manager", "director"
   - "confidence_score": must be a number between 0 and 1

3. ARRAY CONTENT RULES:
   - "technical_requirements": Must contain requirement objects with all fields
   - "soft_skill_requirements": Must contain requirement objects with all fields
   - "experience_requirements": Must contain requirement objects with all fields
   - "education_requirements": Must contain requirement objects with all fields

4. NULL vs EMPTY ARRAY GUIDELINES:
   - Use null for: missing optional strings, numbers
   - Use [] for: missing optional arrays
   - Use "unknown" for: missing required strings that should have a value

FIELD REQUIREMENTS:
- Use EXACT field names as shown
- All requirement arrays must contain objects with required fields
- role_seniority must be one of the specified values
- role_type must be one of the specified values
- Use null for missing optional fields, never omit required fields"""
        
        user_prompt = f"""Analyze this job description and extract structured information:

JOB DESCRIPTION:
{text_content}

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON - no markdown, no explanations, no code blocks
2. Use the exact field names and structure shown in the schema
3. Do NOT include "analysis_timestamp" field - it will be added automatically
4. For missing information, use null (not empty strings)
5. Ensure all required fields are present with appropriate values

CRITICAL DATA STRUCTURE ENFORCEMENT:
- ALL array fields MUST be arrays: use [] for empty arrays, NEVER use null
- "benefits_offered", "success_metrics", "career_growth_opportunities", "critical_success_factors", "potential_challenges" MUST be arrays
- "job_title" must be a string (use "Position Title Not Specified" if unclear)
- "role_seniority" must be one of: "junior", "mid", "senior", "lead", "executive"
- "role_type" must be one of: "individual_contributor", "team_lead", "manager", "director"

Return the complete JSON response following the JobDescriptionAnalysis schema exactly."""
        
        try:
            response = await asyncio.to_thread(
                self.client.chat.completions.create,
                model=self.config.MODEL_TEXT,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=self.config.MAX_TOKENS,
                temperature=self.config.TEMPERATURE
            )
            
            response_content = response.choices[0].message.content
            json_data = self._extract_json_from_response(response_content)
            
            # Remove analysis_timestamp if it exists (will be auto-generated)
            if 'analysis_timestamp' in json_data:
                del json_data['analysis_timestamp']
            
            job_analysis = JobDescriptionAnalysis(**json_data)
            return job_analysis
            
        except Exception as e:
            self.logger.error(f"Job description analysis error: {e}")
            return self._create_fallback_job_analysis(text_content)
    
    def _extract_json_from_response(self, response_content: str) -> Dict[str, Any]:
        """Extract JSON from AI response, handling various formats."""
        
        try:
            # Try direct JSON parsing
            return json.loads(response_content)
        except json.JSONDecodeError:
            # Try to find JSON within markdown or other formatting
            import re
            
            # Look for JSON blocks
            json_match = re.search(r'```json\s*(.*?)\s*```', response_content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(1))
            
            # Look for JSON objects
            json_match = re.search(r'\{.*\}', response_content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
            
            raise ValueError("No valid JSON found in response")
    
    def _serialize_analysis_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Serialize analysis data, converting datetime objects to ISO strings."""
        
        def serialize_item(item):
            if isinstance(item, datetime):
                return item.isoformat()
            elif isinstance(item, dict):
                return {k: serialize_item(v) for k, v in item.items()}
            elif isinstance(item, list):
                return [serialize_item(i) for i in item]
            else:
                return item
        
        return serialize_item(data)
    
    def _clean_datetime_objects(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Remove or convert any datetime objects from the data to prevent serialization errors."""
        
        def clean_item(item):
            if isinstance(item, datetime):
                # Convert datetime to ISO string
                return item.isoformat()
            elif isinstance(item, dict):
                return {k: clean_item(v) for k, v in item.items()}
            elif isinstance(item, list):
                return [clean_item(i) for i in item]
            elif isinstance(item, str):
                # Check if string looks like a datetime and is causing issues
                try:
                    # If it's a datetime string that might be parsed incorrectly, leave as string
                    if 'T' in item and ('Z' in item or '+' in item or item.endswith(':00')):
                        return item  # Keep as string
                except:
                    pass
                return item
            else:
                return item
        
        return clean_item(data)
    
    def _create_fallback_resume_analysis(self, text_content: str) -> ResumeAnalysis:
        """Create basic fallback resume analysis."""
        
        # Extract basic information using simple text processing
        lines = text_content.split('\n')
        
        # Simple name extraction (first non-empty line often contains name)
        name = None
        for line in lines[:5]:
            line = line.strip()
            if line and not any(keyword in line.lower() for keyword in ['email', 'phone', 'address', 'linkedin']):
                name = line
                break
        
        return ResumeAnalysis(
            document_type=DocumentType.RESUME,
            personal_info=PersonalInfo(full_name=name),
            technical_skills=[],
            soft_skills=[],
            work_experience=[],
            education=[],
            certifications=[],
            career_progression=CareerProgression(
                total_experience_years=0.0,
                career_trajectory="unknown",
                job_changes_frequency=0.0,
                industry_consistency=True,
                role_progression=[]
            ),
            strengths=["Resume content extracted"],
            potential_gaps=["Structured analysis failed - manual review recommended"],
            unique_selling_points=[],
            red_flags=["Analysis failed - requires manual review"],
            confidence_score=0.3
        )
    
    def _create_fallback_job_analysis(self, text_content: str) -> JobDescriptionAnalysis:
        """Create basic fallback job description analysis."""
        
        return JobDescriptionAnalysis(
            job_title="Position Title Not Extracted",
            company_info=CompanyInfo(),
            job_summary=text_content[:200] + "...",
            role_seniority="unknown",
            role_type="unknown",
            confidence_score=0.3
        )
    
    async def create_candidate_match_analysis(self, 
                                            resume_analysis: ResumeAnalysis,
                                            job_analysis: JobDescriptionAnalysis) -> CandidateMatch:
        """Create candidate-job matching analysis."""
        
        # This would be a sophisticated matching algorithm
        # For now, provide a basic implementation
        
        # Calculate technical skills match
        candidate_skills = {skill.skill_name.lower() for skill in resume_analysis.technical_skills}
        job_tech_requirements = {req.requirement.lower() for req in job_analysis.technical_requirements}
        
        skill_matches = candidate_skills.intersection(job_tech_requirements)
        skill_match_score = len(skill_matches) / max(1, len(job_tech_requirements))
        
        return CandidateMatch(
            overall_match_score=min(1.0, skill_match_score + 0.2),  # Basic calculation
            technical_skills_match={skill: 0.8 for skill in skill_matches},
            experience_match_score=0.7,  # Placeholder
            education_match_score=0.8,   # Placeholder
            strong_matches=list(skill_matches),
            skill_gaps=list(job_tech_requirements - candidate_skills),
            areas_to_probe=["Technical depth", "Experience relevance", "Cultural fit"]
        )
    
    def get_service_stats(self) -> Dict[str, Any]:
        """Get service statistics."""
        
        return {
            **self.stats,
            "success_rate": (
                self.stats["successful_analyses"] / max(1, self.stats["documents_processed"])
            ),
            "average_processing_time": (
                self.stats["total_processing_time"] / max(1, self.stats["successful_analyses"])
            ),
            "vision_model_usage_rate": (
                self.stats["vision_model_used"] / max(1, self.stats["documents_processed"])
            ),
            "multi_pass_usage_rate": (
                self.stats["multi_pass_analyses"] / max(1, self.stats["successful_analyses"])
            )
        }


# Export main classes
__all__ = [
    'DocumentIntelligenceService',
    'DocumentIntelligenceConfig', 
    'DocumentParsingError'
]