"""
Document Parser Utilities
Handles parsing of different document formats (PDF, DOCX, TXT) for resume and job description analysis.
"""

import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime

try:
    import PyPDF2
    import docx
    from docx import Document
except ImportError as e:
    print(f"Document Parser: Missing required dependency: {e}")
    print("Please install dependencies: pip install PyPDF2 python-docx")
    # Continue without these imports for now


class DocumentParser:
    """Parser for multiple document formats."""
    
    def __init__(self):
        self.logger = self._setup_logger()
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        
    def _setup_logger(self) -> logging.Logger:
        """Set up logging for document parser."""
        logger = logging.getLogger("DocumentParser")
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - Document Parser: %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing parsed content and metadata
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            self.logger.info(f"Parsing document: {file_path.name}")
            
            # Parse based on file type
            if file_extension == '.pdf':
                content = self._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                content = self._parse_docx(file_path)
            elif file_extension == '.txt':
                content = self._parse_txt(file_path)
            else:
                raise ValueError(f"Unsupported format: {file_extension}")
            
            # Extract metadata
            metadata = self._extract_metadata(file_path)
            
            result = {
                "success": True,
                "content": content,
                "metadata": metadata,
                "file_path": str(file_path),
                "parsing_timestamp": datetime.now().isoformat()
            }
            
            self.logger.info(f"Successfully parsed {file_path.name} ({len(content)} characters)")
            return result
            
        except Exception as e:
            self.logger.error(f"Failed to parse {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "file_path": str(file_path),
                "parsing_timestamp": datetime.now().isoformat()
            }
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF document and extract text."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                
                return "\n\n".join(text_content)
                
        except Exception as e:
            self.logger.error(f"PDF parsing error: {e}")
            raise
    
    def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX document and extract text."""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            self.logger.error(f"DOCX parsing error: {e}")
            raise
    
    def _parse_txt(self, file_path: Path) -> str:
        """Parse plain text document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                self.logger.error(f"Text file parsing error: {e}")
                raise
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from document file."""
        try:
            stat = file_path.stat()
            return {
                "file_name": file_path.name,
                "file_size": stat.st_size,
                "file_extension": file_path.suffix.lower(),
                "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "file_path": str(file_path)
            }
        except Exception as e:
            self.logger.warning(f"Failed to extract metadata: {e}")
            return {"file_name": file_path.name, "file_extension": file_path.suffix.lower()}
    
    def validate_document(self, file_path: str) -> bool:
        """
        Validate if document can be parsed.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            True if document is valid, False otherwise
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                self.logger.error(f"Document not found: {file_path}")
                return False
            
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                self.logger.error(f"Unsupported format: {file_extension}")
                return False
            
            # Check file size (max 50MB)
            if file_path.stat().st_size > 50 * 1024 * 1024:
                self.logger.error(f"File too large: {file_path.stat().st_size} bytes")
                return False
            
            return True
            
        except Exception as e:
            self.logger.error(f"Document validation error: {e}")
            return False
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats."""
        return self.supported_formats.copy()
    
    def extract_sections(self, content: str) -> Dict[str, str]:
        """
        Extract common sections from document content.
        
        Args:
            content: Raw document text content
            
        Returns:
            Dict mapping section names to content
        """
        sections = {}
        
        # Common section headers
        section_headers = [
            "experience", "education", "skills", "summary", "objective",
            "work history", "employment", "qualifications", "achievements",
            "projects", "certifications", "languages", "interests",
            "requirements", "responsibilities", "duties", "benefits"
        ]
        
        lines = content.split('\n')
        current_section = "general"
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            is_header = False
            for header in section_headers:
                if header in line_lower and len(line.strip()) < 100:
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    
                    # Start new section
                    current_section = header
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections


# Convenience function for quick parsing
def parse_document(file_path: str) -> Dict[str, Any]:
    """Quick function to parse a document."""
    parser = DocumentParser()
    return parser.parse_document(file_path)


# Export main class
__all__ = ['DocumentParser', 'parse_document']
